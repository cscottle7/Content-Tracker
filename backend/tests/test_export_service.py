"""Tests for export_service module.

Tests DOCX and PDF export functionality.
"""

import pytest
import os
from pathlib import Path
from datetime import date
from docx import Document

from app.services import export_service
from app.models.content import ContentResponse


@pytest.fixture
def sample_content_items():
    """Sample content items for export testing."""
    return [
        ContentResponse(
            id="test-1",
            file_path="/tmp/test1.md",
            title="Blog Post: SEO Best Practices",
            content_type="blog",
            status="published",
            created_date=date(2024, 1, 15),
            updated_date=date(2024, 1, 20),
            publish_date=date(2024, 1, 16),
            author="John Doe",
            url="https://example.com/seo-best-practices",
            description="A comprehensive guide to SEO optimization.",
            tags=["seo", "marketing", "tutorial"],
            categories=["Content Marketing"],
            custom_fields={"client": "ClientA"},
            body="# SEO Best Practices\n\nSearch Engine Optimization is crucial..."
        ),
        ContentResponse(
            id="test-2",
            file_path="/tmp/test2.md",
            title="Video: Product Demo",
            content_type="video",
            status="published",
            created_date=date(2024, 1, 18),
            updated_date=date(2024, 1, 22),
            publish_date=date(2024, 1, 20),
            author="Jane Smith",
            url="https://youtube.com/watch?v=abc123",
            description="Product demonstration video for new feature.",
            tags=["product", "demo", "video"],
            categories=["Product Marketing"],
            custom_fields={"client": "ClientB"},
            body="Video script and talking points..."
        ),
        ContentResponse(
            id="test-3",
            file_path="/tmp/test3.md",
            title="Podcast Episode: Industry Trends",
            content_type="podcast",
            status="draft",
            created_date=date(2024, 1, 25),
            updated_date=date(2024, 1, 25),
            author="Sarah Johnson",
            description="Discussion of current industry trends.",
            tags=["podcast", "trends"],
            categories=["Thought Leadership"],
            custom_fields={},
            body="Episode notes and timestamps..."
        ),
    ]


@pytest.mark.asyncio
async def test_export_to_docx(sample_content_items, tmp_path, monkeypatch):
    """Test DOCX export generation."""
    # Set temporary export path
    monkeypatch.setattr("app.services.export_service.settings.EXPORTS_PATH", str(tmp_path))

    # Export to DOCX
    file_path = await export_service.export_to_docx(
        content_items=sample_content_items,
        title="Test Export Report",
        include_fields=["title", "content_type", "status", "author", "tags"]
    )

    # Verify file created
    assert os.path.exists(file_path)
    assert file_path.endswith(".docx")

    # Verify DOCX structure
    doc = Document(file_path)

    # Check title is in document
    assert "Test Export Report" in doc.paragraphs[0].text

    # Check summary includes item count
    text_content = "\n".join([p.text for p in doc.paragraphs])
    assert "Total Items: 3" in text_content
    assert "Blog Post: SEO Best Practices" in text_content
    assert "Video: Product Demo" in text_content
    assert "Podcast Episode: Industry Trends" in text_content


@pytest.mark.asyncio
async def test_export_to_docx_with_all_fields(sample_content_items, tmp_path, monkeypatch):
    """Test DOCX export with all fields included."""
    monkeypatch.setattr("app.services.export_service.settings.EXPORTS_PATH", str(tmp_path))

    file_path = await export_service.export_to_docx(
        content_items=sample_content_items,
        title="Complete Export",
        include_fields=None  # Include all fields
    )

    assert os.path.exists(file_path)

    doc = Document(file_path)

    # Get all text from paragraphs and tables
    all_text = []
    for paragraph in doc.paragraphs:
        all_text.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text.append(cell.text)

    text_content = "\n".join(all_text)

    # Verify detailed fields are present
    assert "John Doe" in text_content
    assert "https://example.com/seo-best-practices" in text_content
    assert "seo, marketing, tutorial" in text_content


@pytest.mark.asyncio
async def test_export_to_pdf(sample_content_items, tmp_path, monkeypatch):
    """Test PDF export generation."""
    monkeypatch.setattr("app.services.export_service.settings.EXPORTS_PATH", str(tmp_path))

    # Export to PDF
    file_path = await export_service.export_to_pdf(
        content_items=sample_content_items,
        title="Test PDF Report",
        include_fields=["title", "content_type", "status", "description"]
    )

    # Verify file created (may be .pdf or .html depending on WeasyPrint availability)
    assert os.path.exists(file_path)
    assert file_path.endswith(".pdf") or file_path.endswith(".html")

    # Verify file is not empty
    file_size = os.path.getsize(file_path)
    assert file_size > 500  # File should have content


@pytest.mark.asyncio
async def test_export_empty_content_list(tmp_path, monkeypatch):
    """Test export with empty content list."""
    monkeypatch.setattr("app.services.export_service.settings.EXPORTS_PATH", str(tmp_path))

    # Export empty list to DOCX
    file_path = await export_service.export_to_docx(
        content_items=[],
        title="Empty Report"
    )

    assert os.path.exists(file_path)

    doc = Document(file_path)
    text_content = "\n".join([p.text for p in doc.paragraphs])

    assert "Total Items: 0" in text_content


@pytest.mark.asyncio
async def test_export_with_custom_template_fallback(sample_content_items, tmp_path, monkeypatch):
    """Test PDF export falls back to default template when custom doesn't exist."""
    monkeypatch.setattr("app.services.export_service.settings.EXPORTS_PATH", str(tmp_path))

    # Request non-existent template
    file_path = await export_service.export_to_pdf(
        content_items=sample_content_items,
        title="Custom Template Test",
        template_name="nonexistent_template"
    )

    # Should still generate using default template (may be .pdf or .html)
    assert os.path.exists(file_path)
    assert file_path.endswith(".pdf") or file_path.endswith(".html")


@pytest.mark.asyncio
async def test_cleanup_old_exports(tmp_path, monkeypatch):
    """Test cleanup of old export files."""
    monkeypatch.setattr("app.services.export_service.settings.EXPORTS_PATH", str(tmp_path))

    # Create test export files
    old_file1 = tmp_path / "old_export.docx"
    old_file2 = tmp_path / "old_export.pdf"
    old_file1.touch()
    old_file2.touch()

    # Modify timestamps to simulate old files (older than 1 hour)
    import time
    old_time = time.time() - (2 * 3600)  # 2 hours ago
    os.utime(old_file1, (old_time, old_time))
    os.utime(old_file2, (old_time, old_time))

    # Create recent file
    recent_file = tmp_path / "recent_export.docx"
    recent_file.touch()

    # Run cleanup (max age 1 hour)
    deleted_count = await export_service.cleanup_old_exports(max_age_hours=1)

    # Verify old files deleted, recent file kept
    assert deleted_count == 2
    assert not old_file1.exists()
    assert not old_file2.exists()
    assert recent_file.exists()


@pytest.mark.asyncio
async def test_export_type_counts(sample_content_items, tmp_path, monkeypatch):
    """Test that export includes correct content type counts."""
    monkeypatch.setattr("app.services.export_service.settings.EXPORTS_PATH", str(tmp_path))

    file_path = await export_service.export_to_docx(
        content_items=sample_content_items,
        title="Type Counts Test"
    )

    doc = Document(file_path)
    text_content = "\n".join([p.text for p in doc.paragraphs])

    # Should show counts by type
    assert "Blog: 1" in text_content or "blog: 1" in text_content.lower()
    assert "Video: 1" in text_content or "video: 1" in text_content.lower()
    assert "Podcast: 1" in text_content or "podcast: 1" in text_content.lower()
