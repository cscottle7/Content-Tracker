"""Export service for generating DOCX and PDF reports.

This module handles the generation of professional reports from content items,
supporting both DOCX (via python-docx) and PDF (via WeasyPrint) formats.
"""

import os
import json
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Dict
from uuid import uuid4

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    # WeasyPrint requires system libraries not available on all platforms
    WEASYPRINT_AVAILABLE = False
    print("Warning: WeasyPrint not available. PDF export will use alternative method.")

from jinja2 import Template

from app.config import settings
from app.models.content import ContentResponse


def _get_export_path(export_id: str, format: str) -> str:
    """
    Get file path for export file.

    Args:
        export_id: Unique identifier for export
        format: File format (docx or pdf)

    Returns:
        Absolute path to export file
    """
    export_dir = Path(settings.EXPORTS_PATH)
    export_dir.mkdir(parents=True, exist_ok=True)
    return str(export_dir / f"{export_id}.{format}")


def _apply_docx_styling(doc: Document) -> None:
    """
    Apply professional styling to DOCX document.

    Args:
        doc: python-docx Document object
    """
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Style headings
    heading1 = doc.styles['Heading 1']
    heading1.font.size = Pt(18)
    heading1.font.color.rgb = RGBColor(31, 78, 121)
    heading1.font.bold = True

    heading2 = doc.styles['Heading 2']
    heading2.font.size = Pt(14)
    heading2.font.color.rgb = RGBColor(68, 114, 196)
    heading2.font.bold = True


async def export_to_docx(
    content_items: List[ContentResponse],
    title: str = "Content Report",
    include_fields: Optional[List[str]] = None,
    template_name: Optional[str] = None
) -> str:
    """
    Export content items to DOCX format.

    Args:
        content_items: List of content items to export
        title: Report title
        include_fields: List of metadata fields to include (None = all)
        template_name: Optional custom template name

    Returns:
        Path to generated DOCX file
    """
    export_id = str(uuid4())
    file_path = _get_export_path(export_id, "docx")

    # Create document
    doc = Document()
    _apply_docx_styling(doc)

    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add generation date
    date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacing

    # Add summary
    summary = doc.add_heading("Summary", level=1)
    doc.add_paragraph(f"Total Items: {len(content_items)}")

    # Count by content type
    type_counts: Dict[str, int] = {}
    for item in content_items:
        type_counts[item.content_type] = type_counts.get(item.content_type, 0) + 1

    doc.add_paragraph("Content by Type:")
    for content_type, count in sorted(type_counts.items()):
        doc.add_paragraph(f"  • {content_type.title()}: {count}", style='List Bullet')

    doc.add_page_break()

    # Add content items
    doc.add_heading("Content Items", level=1)

    # Default fields to include
    if include_fields is None:
        include_fields = [
            'title', 'content_type', 'status', 'author', 'publish_date',
            'url', 'description', 'tags', 'categories', 'created_date', 'updated_date'
        ]

    for idx, item in enumerate(content_items, 1):
        # Item heading
        doc.add_heading(f"{idx}. {item.title}", level=2)

        # Metadata table
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light Grid Accent 1'

        # Add fields
        if 'content_type' in include_fields:
            row = table.add_row()
            row.cells[0].text = "Type"
            row.cells[1].text = item.content_type.title()

        if 'status' in include_fields:
            row = table.add_row()
            row.cells[0].text = "Status"
            row.cells[1].text = item.status.title()

        if 'author' in include_fields and item.author:
            row = table.add_row()
            row.cells[0].text = "Author"
            row.cells[1].text = item.author

        if 'publish_date' in include_fields and item.publish_date:
            row = table.add_row()
            row.cells[0].text = "Publish Date"
            row.cells[1].text = item.publish_date.strftime('%Y-%m-%d')

        if 'url' in include_fields and item.url:
            row = table.add_row()
            row.cells[0].text = "URL"
            row.cells[1].text = item.url

        if 'created_date' in include_fields:
            row = table.add_row()
            row.cells[0].text = "Created"
            row.cells[1].text = item.created_date.strftime('%Y-%m-%d')

        if 'updated_date' in include_fields:
            row = table.add_row()
            row.cells[0].text = "Updated"
            row.cells[1].text = item.updated_date.strftime('%Y-%m-%d')

        # Description
        if 'description' in include_fields and item.description:
            doc.add_paragraph()
            doc.add_paragraph("Description:", style='Heading 3')
            doc.add_paragraph(item.description)

        # Tags and categories
        if 'tags' in include_fields and item.tags:
            doc.add_paragraph()
            doc.add_paragraph(f"Tags: {', '.join(item.tags)}")

        if 'categories' in include_fields and item.categories:
            doc.add_paragraph(f"Categories: {', '.join(item.categories)}")

        # Content body (if available)
        if item.body and 'body' in include_fields:
            doc.add_paragraph()
            doc.add_paragraph("Content:", style='Heading 3')
            # Add body as plain text (markdown not rendered in DOCX)
            doc.add_paragraph(item.body)

        # Spacing between items
        if idx < len(content_items):
            doc.add_paragraph()
            doc.add_paragraph("─" * 80)
            doc.add_paragraph()

    # Save document
    doc.save(file_path)

    return file_path


async def export_to_pdf(
    content_items: List[ContentResponse],
    title: str = "Content Report",
    include_fields: Optional[List[str]] = None,
    template_name: Optional[str] = None
) -> str:
    """
    Export content items to PDF format.

    Args:
        content_items: List of content items to export
        title: Report title
        include_fields: List of metadata fields to include (None = all)
        template_name: Optional custom template name

    Returns:
        Path to generated PDF file
    """
    export_id = str(uuid4())
    file_path = _get_export_path(export_id, "pdf")

    # Default fields to include
    if include_fields is None:
        include_fields = [
            'title', 'content_type', 'status', 'author', 'publish_date',
            'url', 'description', 'tags', 'categories', 'created_date', 'updated_date'
        ]

    # Count by content type
    type_counts: Dict[str, int] = {}
    for item in content_items:
        type_counts[item.content_type] = type_counts.get(item.content_type, 0) + 1

    # Load template (or use default)
    template_path = Path(settings.EXPORTS_PATH) / "templates" / "default.html"

    if template_name and template_name != "default":
        custom_template = Path(settings.EXPORTS_PATH) / "templates" / f"{template_name}.html"
        if custom_template.exists():
            template_path = custom_template

    # Check if template exists, otherwise use inline template
    if template_path.exists():
        with open(template_path, 'r', encoding='utf-8') as f:
            html_template = f.read()
    else:
        # Default inline HTML template
        html_template = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ title }}</title>
    <style>
        @page {
            size: A4;
            margin: 2cm;
            @bottom-right {
                content: "Page " counter(page) " of " counter(pages);
                font-size: 9pt;
                color: #666;
            }
        }
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
        }
        h1 {
            color: #1f4e79;
            border-bottom: 3px solid #4472c4;
            padding-bottom: 10px;
            text-align: center;
            font-size: 28pt;
        }
        h2 {
            color: #4472c4;
            margin-top: 30px;
            font-size: 18pt;
            page-break-after: avoid;
        }
        h3 {
            color: #5b9bd5;
            margin-top: 20px;
            font-size: 14pt;
        }
        .meta-info {
            text-align: center;
            color: #666;
            margin-bottom: 30px;
        }
        .summary {
            background: #f0f4f8;
            padding: 20px;
            border-left: 4px solid #4472c4;
            margin-bottom: 30px;
        }
        .summary ul {
            list-style: none;
            padding-left: 0;
        }
        .summary li {
            padding: 5px 0;
        }
        .content-item {
            margin-bottom: 40px;
            page-break-inside: avoid;
        }
        .metadata-table {
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }
        .metadata-table th {
            background: #4472c4;
            color: white;
            padding: 10px;
            text-align: left;
            width: 30%;
        }
        .metadata-table td {
            padding: 10px;
            border: 1px solid #ddd;
        }
        .metadata-table tr:nth-child(even) {
            background: #f9f9f9;
        }
        .tags, .categories {
            display: inline-block;
            background: #e7f3ff;
            color: #0066cc;
            padding: 3px 10px;
            border-radius: 3px;
            margin: 2px;
            font-size: 9pt;
        }
        .description {
            background: #f9f9f9;
            padding: 15px;
            border-left: 3px solid #5b9bd5;
            margin: 15px 0;
        }
        .divider {
            border-top: 2px dashed #ccc;
            margin: 30px 0;
        }
    </style>
</head>
<body>
    <h1>{{ title }}</h1>
    <div class="meta-info">Generated: {{ generation_date }}</div>

    <div class="summary">
        <h2>Summary</h2>
        <p><strong>Total Items:</strong> {{ total_items }}</p>
        <p><strong>Content by Type:</strong></p>
        <ul>
        {% for content_type, count in type_counts.items() %}
            <li>{{ content_type|title }}: {{ count }}</li>
        {% endfor %}
        </ul>
    </div>

    <h2>Content Items</h2>

    {% for item in content_items %}
    <div class="content-item">
        <h2>{{ loop.index }}. {{ item.title }}</h2>

        <table class="metadata-table">
            {% if 'content_type' in include_fields %}
            <tr>
                <th>Type</th>
                <td>{{ item.content_type|title }}</td>
            </tr>
            {% endif %}

            {% if 'status' in include_fields %}
            <tr>
                <th>Status</th>
                <td>{{ item.status|title }}</td>
            </tr>
            {% endif %}

            {% if 'author' in include_fields and item.author %}
            <tr>
                <th>Author</th>
                <td>{{ item.author }}</td>
            </tr>
            {% endif %}

            {% if 'publish_date' in include_fields and item.publish_date %}
            <tr>
                <th>Publish Date</th>
                <td>{{ item.publish_date }}</td>
            </tr>
            {% endif %}

            {% if 'url' in include_fields and item.url %}
            <tr>
                <th>URL</th>
                <td><a href="{{ item.url }}">{{ item.url }}</a></td>
            </tr>
            {% endif %}

            {% if 'created_date' in include_fields %}
            <tr>
                <th>Created</th>
                <td>{{ item.created_date }}</td>
            </tr>
            {% endif %}

            {% if 'updated_date' in include_fields %}
            <tr>
                <th>Updated</th>
                <td>{{ item.updated_date }}</td>
            </tr>
            {% endif %}
        </table>

        {% if 'description' in include_fields and item.description %}
        <div class="description">
            <h3>Description</h3>
            <p>{{ item.description }}</p>
        </div>
        {% endif %}

        {% if 'tags' in include_fields and item.tags %}
        <p>
            <strong>Tags:</strong>
            {% for tag in item.tags %}
            <span class="tags">{{ tag }}</span>
            {% endfor %}
        </p>
        {% endif %}

        {% if 'categories' in include_fields and item.categories %}
        <p>
            <strong>Categories:</strong>
            {% for category in item.categories %}
            <span class="categories">{{ category }}</span>
            {% endfor %}
        </p>
        {% endif %}

        {% if item.body and 'body' in include_fields %}
        <div class="description">
            <h3>Content</h3>
            <pre style="white-space: pre-wrap; font-family: inherit;">{{ item.body }}</pre>
        </div>
        {% endif %}
    </div>

    {% if not loop.last %}
    <div class="divider"></div>
    {% endif %}
    {% endfor %}
</body>
</html>
        """

    # Render template
    template = Template(html_template)
    html_content = template.render(
        title=title,
        generation_date=datetime.now().strftime('%Y-%m-%d %H:%M'),
        total_items=len(content_items),
        type_counts=type_counts,
        content_items=[
            {
                'title': item.title,
                'content_type': item.content_type,
                'status': item.status,
                'author': item.author,
                'publish_date': item.publish_date.strftime('%Y-%m-%d') if item.publish_date else None,
                'url': item.url,
                'created_date': item.created_date.strftime('%Y-%m-%d'),
                'updated_date': item.updated_date.strftime('%Y-%m-%d'),
                'description': item.description,
                'tags': item.tags,
                'categories': item.categories,
                'body': item.body,
            }
            for item in content_items
        ],
        include_fields=include_fields
    )

    # Generate PDF
    if WEASYPRINT_AVAILABLE:
        HTML(string=html_content).write_pdf(file_path)
    else:
        # Fallback: Save as HTML with note that PDF requires server setup
        html_file_path = file_path.replace('.pdf', '.html')
        with open(html_file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        # Return HTML path instead
        return html_file_path

    return file_path


async def cleanup_old_exports(max_age_hours: int = 1) -> int:
    """
    Clean up export files older than specified age.

    Args:
        max_age_hours: Maximum age of export files in hours

    Returns:
        Number of files deleted
    """
    export_dir = Path(settings.EXPORTS_PATH)
    if not export_dir.exists():
        return 0

    current_time = datetime.now()
    deleted_count = 0

    # Combine all export files (docx, pdf, html)
    all_files = list(export_dir.glob("*.docx")) + list(export_dir.glob("*.pdf")) + list(export_dir.glob("*.html"))

    for file_path in all_files:
        # Get file modification time
        file_mtime = datetime.fromtimestamp(file_path.stat().st_mtime)
        age_hours = (current_time - file_mtime).total_seconds() / 3600

        if age_hours > max_age_hours:
            file_path.unlink()
            deleted_count += 1

    return deleted_count
